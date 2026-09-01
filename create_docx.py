import sys
from docx import Document

def txt_to_docx(txt_file, docx_file):
    doc = Document()
    
    with open(txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for idx, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # The first line is the Title
        if idx == 0:
            doc.add_heading(line, 0)
        # Check if line starts with a number and a dot like "1. " or "5. "
        elif len(line) > 2 and line[0].isdigit() and line[1] == '.':
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_file)
    print(f"Successfully saved {docx_file}")

if __name__ == "__main__":
    txt_to_docx("hackathon_report.txt", "hackathon_report.docx")
